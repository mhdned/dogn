# requirments
import os

# PyMuPDF
import fitz
from docx import Document
from docx.shared import Inches
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from PIL import Image

from configs.config import config

# instance
router = APIRouter()


upload_directory = f"{config["PATH"]}{config["UPLOAD_PATH"]}/document"

# routes
@router.post("/sign")
async def sign_document(
    file: UploadFile = File(...),
    signature: UploadFile = File(...),
    position_x: int = Form(...),
    position_y: int = Form(...),
):
    # split extention from uploaded file (on document)
    file_extension = file.filename.split(".")[-1].lower()

    # check type of document
    if file_extension not in ["pdf", "docx"]:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Only PDF and DOCX are allowed.",
        )

    # Save uploaded files temporarily
    input_file_path = f"{upload_directory}/temp/{file.filename}"
    signature_file_path = f"{upload_directory}/temp/{signature.filename}"
    
    
    # store files in upload folder
    with open(input_file_path, "wb") as f:
        f.write(await file.read())

    with open(signature_file_path, "wb") as f:
        f.write(await signature.read())

    # create result of sign document
    output_file_path = f"{upload_directory}/test/signed_{file.filename}"

    # pdf sign
    if file_extension == "pdf":
        append_signature_to_pdf(
            input_file_path,
            signature_file_path,
            output_file_path,
            position_x,
            position_y,
        )
    # word sign
    elif file_extension == "docx":
        append_signature_to_docx(input_file_path, signature_file_path, output_file_path)

    return FileResponse(output_file_path, filename=f"signed_{file.filename}")


def append_signature_to_pdf(
    input_pdf_path: str,
    signature_image_path: str,
    output_pdf_path: str,
    position_x: int,
    position_y: int,
):
    pdf_document = fitz.open(input_pdf_path)
    signature_image = Image.open(signature_image_path)

    signature_image_path_temp = f"{upload_directory}/temp/temp_signature.png"
    signature_image.save(signature_image_path_temp)

    for page in pdf_document:
        img_rect = fitz.Rect(
            position_x,
            position_y,
            position_x + signature_image.width,
            position_y + signature_image.height,
        )
        page.insert_image(img_rect, filename=signature_image_path_temp)

    pdf_document.save(output_pdf_path)


def append_signature_to_docx(
    input_docx_path: str, signature_image_path: str, output_docx_path: str
):
    document = Document(input_docx_path)
    signature_image = Image.open(signature_image_path)

    last_paragraph = document.paragraphs[-1]
    run = last_paragraph.add_run()
    run.add_picture(signature_image, width=Inches(2))

    document.save(output_docx_path)
